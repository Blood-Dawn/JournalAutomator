import os
import sys

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

import pytest
from journal_updater import journal_updater as ju


def test_apply_page_borders(tmp_path):
    doc = ju.Document()
    border_specs = {
        "left": {"val": "single", "sz": 4, "space": 0, "color": "000000"},
        "right": {"val": "single", "sz": 4, "space": 0, "color": "000000"},
    }
    ju.apply_page_borders(doc, 0, border_specs)

    from docx.oxml.ns import qn

    borders = list(doc.sections[0]._sectPr.findall(qn("w:pgBorders")))
    assert len(borders) == 1


def test_validate_issue_number_and_volume():
    doc = ju.Document()
    doc.add_paragraph("Volume 5, Issue 2")
    doc.add_paragraph("2024")
    ju.validate_issue_number_and_volume(doc, "5", "2", "2024")

    with pytest.raises(ValueError):
        ju.validate_issue_number_and_volume(doc, "1", "2", "2024")
