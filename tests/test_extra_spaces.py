import sys
import os
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

import journal_updater.journal_updater as ju
from docx.enum.text import WD_BREAK


def _build_doc():
    doc = ju.Document()
    p1 = doc.add_paragraph("Page 1  text")
    br = p1.add_run()
    br.add_break(WD_BREAK.PAGE)
    doc.add_paragraph("Page 2  text")
    return doc


def test_detect_and_remove_extra_spaces_page_range():
    doc = _build_doc()
    ju.detect_and_remove_extra_spaces(doc, [1])
    assert doc.paragraphs[0].text == "Page 1 text"
    assert doc.paragraphs[1].text == "Page 2  text"

    ju.detect_and_remove_extra_spaces(doc, [2])
    assert doc.paragraphs[1].text == "Page 2 text"

def test_map_pages_to_paragraphs():
    doc = _build_doc()
    pages = ju.map_pages_to_paragraphs(doc)
    assert 1 in pages and 2 in pages
    assert any(p.text.startswith("Page 1") for p in pages[1])
    assert any(p.text.startswith("Page 2") for p in pages[2])
