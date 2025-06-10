import os
import sys
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

import json
from docx import Document
from docx.enum.text import WD_BREAK
import journal_updater.journal_updater as ju


def test_autofit_first_table_helper():
    doc = Document()
    p = doc.add_paragraph("page1")
    p.add_run().add_break(WD_BREAK.PAGE)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    ju.autofit_first_table(doc, 2)
    assert tbl.autofit is True


def test_update_journal_autofit_instruction(tmp_path):
    base = Document()
    base.add_paragraph("Volume 1, Issue 1")
    p = base.add_paragraph("intro")
    p.add_run().add_break(WD_BREAK.PAGE)
    tbl = base.add_table(rows=1, cols=2)
    tbl.cell(0, 0).text = "target"
    tbl.autofit = False
    base.add_paragraph("ARTICLES")
    base_path = tmp_path / "base.docx"
    base.save(base_path)

    content = tmp_path / "content"
    content.mkdir()
    (content / "instructions.json").write_text(json.dumps({"autofit_table_on_page": 2}))

    out_path = tmp_path / "out.docx"
    ju.update_journal(base_path, content, out_path, "1", "1", "June 2025", article_files=[])
    result = Document(out_path)
    target_table = next(t for t in result.tables if t.cell(0, 0).text == "target")
    assert target_table.autofit is True
