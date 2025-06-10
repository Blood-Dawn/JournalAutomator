import sys
import os

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
import journal_updater.journal_updater as ju
from docx.enum.text import WD_BREAK


def test_delete_after_page(tmp_path):
    base = ju.Document()
    first = base.add_paragraph("Keep this")
    first.add_run().add_break(WD_BREAK.PAGE)
    base.add_paragraph("Remove")
    ju.add_page_borders(base, 0)

    base_path = tmp_path / "base.docx"
    base.save(base_path)

    content = tmp_path / "content"
    content.mkdir()

    out_path = tmp_path / "out.docx"
    ju.update_journal(
        base_path,
        content,
        out_path,
        volume="1",
        issue="1",
        month_year="June 2025",
        start_page=1,
    )

    result = ju.Document(out_path)
    texts = [p.text for p in result.paragraphs]
    assert texts == ["Keep this"]

    from docx.oxml.ns import qn

    assert not list(result.sections[0]._sectPr.findall(qn("w:pgBorders")))
