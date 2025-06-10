import sys
import os
import json

from docx.shared import Pt
from docx.enum.text import WD_BREAK

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
import journal_updater.journal_updater as journal_updater


def test_replace_text():
    class FakeRun:
        def __init__(self, text):
            self.text = text

    class FakeParagraph:
        def __init__(self, text):
            self.text = text
            self.runs = [FakeRun(text)]

    paragraphs = [FakeParagraph("2023 subscription info"), FakeParagraph("nothing")]
    journal_updater.replace_text_in_paragraphs(paragraphs, "2023", "")
    assert paragraphs[0].runs[0].text == " subscription info"

def test_load_and_update_front_cover(tmp_path):
    doc_path = tmp_path / "base.docx"
    doc = journal_updater.Document()
    p = doc.add_paragraph("Volume 1, Issue 3")
    doc.save(doc_path)

    loaded = journal_updater.load_document(doc_path)
    journal_updater.update_front_cover(loaded, "1", "1", "June 2025", 1)
    assert "June 2025" in loaded.paragraphs[0].text
    assert "Page" not in loaded.paragraphs[0].text
    assert all(run.font.bold for run in loaded.paragraphs[0].runs)


def test_update_page2_header(tmp_path):
    doc = journal_updater.Document()
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.text = "Old header"
    journal_updater.update_page2_header(doc, "Volume 1", 2)
    assert "Page 2" in section.header.paragraphs[0].text


def test_font_utils():
    doc = journal_updater.Document()
    doc.add_paragraph("p0")
    doc.add_paragraph("p1")

    journal_updater.set_font_size(doc, 1, 16)
    journal_updater.set_line_spacing(doc, 1, 1.5)

    assert doc.paragraphs[1].runs[0].font.size.pt == 16
    assert doc.paragraphs[1].paragraph_format.line_spacing == 1.5


def test_update_journal_formatting(tmp_path):
    base = journal_updater.Document()
    base.add_paragraph("ARTICLES")
    base_path = tmp_path / "base.docx"
    base.save(base_path)

    content_dir = tmp_path / "content"
    content_dir.mkdir()
    art = journal_updater.Document()
    art.add_paragraph("Article text")
    art.save(content_dir / "article1.docx")

    import json

    (content_dir / "instructions.json").write_text(
        json.dumps({
            "font_size": 14,
            "line_spacing": 2,
            "font_family": "Times New Roman",
        })
    )

    out_path = tmp_path / "out.docx"
    journal_updater.update_journal(
        base_path,
        content_dir,
        out_path,
        "1",
        "1",
        "June 2025",
    )
    result = journal_updater.Document(out_path)
    first = result.sections[0]
    assert first.different_first_page_header_footer
    assert len(first.first_page_footer.tables) == 0
    assert len(first.footer.tables) == 1

    assert result.paragraphs[0].runs[0].font.size.pt == 14
    assert result.paragraphs[0].paragraph_format.line_spacing == 2
    assert result.paragraphs[0].runs[0].font.name == "Times New Roman"

def test_format_front_and_footer(tmp_path):
    doc = journal_updater.Document()
    p = doc.add_paragraph("Volume 1")
    footer_p = doc.sections[0].footer.paragraphs[0]
    footer_p.text = "footer"

    journal_updater.format_front_and_footer(doc, font_size=13, line_spacing=1.25)
    assert doc.paragraphs[0].runs[0].font.size.pt == 13
    assert doc.paragraphs[0].paragraph_format.line_spacing == 1.25
    assert doc.sections[0].footer.paragraphs[0].runs[0].font.size.pt == 13
    assert doc.sections[0].footer.paragraphs[0].paragraph_format.line_spacing == 1.25


def test_set_font_size_and_spacing_from_page():
    doc = journal_updater.Document()
    p1 = doc.add_paragraph("p1")
    p1.runs[0].font.size = Pt(10)
    p1.paragraph_format.line_spacing = 1
    br = p1.add_run()
    br.add_break(WD_BREAK.PAGE)
    p2 = doc.add_paragraph("p2")
    p2.runs[0].font.size = Pt(10)
    p2.paragraph_format.line_spacing = 1
    p3 = doc.add_paragraph("p3")
    p3.runs[0].font.size = Pt(10)
    p3.paragraph_format.line_spacing = 1

    journal_updater.set_font_size_from_page(doc, 2, 16)
    journal_updater.set_line_spacing_from_page(doc, 2, 2)

    assert p1.runs[0].font.size.pt == 10
    assert p1.paragraph_format.line_spacing == 1
    assert p2.runs[0].font.size.pt == 16
    assert p3.runs[0].font.size.pt == 16
    assert p2.paragraph_format.line_spacing == 2
    assert p3.paragraph_format.line_spacing == 2


def test_update_journal_page_specific_formatting(tmp_path):
    base = journal_updater.Document()
    p1 = base.add_paragraph("Volume 1, Issue 1")
    base.add_paragraph("").add_run().add_break(WD_BREAK.PAGE)
    base.add_paragraph("ARTICLES")
    base_path = tmp_path / "base.docx"
    base.save(base_path)

    content_dir = tmp_path / "content"
    content_dir.mkdir()
    art = journal_updater.Document()
    art.add_paragraph("Article text")
    art.save(content_dir / "article1.docx")

    import json

    (content_dir / "instructions.json").write_text(
        json.dumps(
            {
                "font_size_from_page": {"page": 2, "size": 14},
                "line_spacing_from_page": {"page": 2, "spacing": 1.5},
            }
        )
    )

    out_path = tmp_path / "out.docx"
    journal_updater.update_journal(
        base_path,
        content_dir,
        out_path,
        "1",
        "1",
        "June 2025",
    )

    result = journal_updater.Document(out_path)
    pages = journal_updater.map_pages_to_paragraphs(result)
    for p in pages.get(2, []):
        assert p.runs[0].font.size.pt == 14
        assert p.paragraph_format.line_spacing == 1.5
