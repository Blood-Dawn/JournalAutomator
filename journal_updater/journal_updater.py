"""Utility functions for updating ABNFF journal Word documents."""

import argparse
import json
import logging
from pathlib import Path
from typing import Dict, Iterable, List, Optional

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def load_document(path: Path) -> Document:
    """Open the Word file at ``path`` and return a ``Document`` object."""
    return Document(str(path))


def save_document(doc: Document, path_out: Path) -> None:
    """Save ``doc`` to ``path_out``."""
    doc.save(str(path_out))


def replace_text_in_paragraphs(paragraphs, search_text, replace_text):
    for p in paragraphs:
        if search_text in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if search_text in inline[i].text:
                    inline[i].text = inline[i].text.replace(search_text, replace_text)


def update_front_cover(
    doc: Document,
    volume: str,
    issue: str,
    month_year: str,
    page_num: int,
) -> None:
    """Update volume/issue block on the front cover."""
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        WD_ALIGN_PARAGRAPH = None  # type: ignore

    search = "Volume"
    for p in doc.paragraphs:
        if search in p.text:
            p.text = f"Volume {volume}, Issue {issue}\n{month_year}"
            for run in p.runs:
                run.font.bold = True
            if WD_ALIGN_PARAGRAPH is not None:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    pPr = p._p.get_or_add_pPr()
                    for b in pPr.findall(qn("w:pBdr")):
                        pPr.remove(b)
                    pBdr = OxmlElement("w:pBdr")
                    bottom = OxmlElement("w:bottom")
                    bottom.set(qn("w:val"), "single")
                    bottom.set(qn("w:sz"), "6")
                    bottom.set(qn("w:space"), "1")
                    bottom.set(qn("w:color"), "000000")
                    pBdr.append(bottom)
                    pPr.append(pBdr)
                except Exception:
                    pass
            break


def update_business_information(
    doc: Document, old_year: str, new_beginning_text: str
) -> None:
    """Update the business information block on page 1."""

    replace_text_in_paragraphs(doc.paragraphs, old_year, "")

    for p in doc.paragraphs:
        if "Annual subscription" in p.text:
            # replace first sentence beginning
            first_part = p.text.split(".")[0]
            if first_part:
                p.text = p.text.replace(first_part, new_beginning_text, 1)
            else:
                p.text = new_beginning_text + p.text
            break


def update_page2_header(doc: Document, new_header_line1: str, page_num: int) -> None:
    """Replace the header text on page 2."""

    for section in doc.sections:
        header = section.header
        text = f"{new_header_line1}\nPage {page_num}"
        for p in header.paragraphs:
            if p.text.strip():
                p.text = text
                break


def format_front_cover(doc: Document) -> None:
    """Bold and center the first paragraph of the document."""
    if not doc.paragraphs:
        return
    p = doc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.bold = True


def layout_footer(doc: Document) -> None:
    """Center footers across all sections."""
    for section in doc.sections:
        footer = section.footer
        paragraph = (
            footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        )
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def update_associate_editors(
    doc: Document, remove_name: str, new_name: str, new_affiliation: str, new_email: str
) -> None:
    """Update the associate editors list on page 2."""

    for p in doc.paragraphs:
        if remove_name in p.text:
            # replace the line with new associate editor info
            p.text = f"{new_name}, {new_affiliation}\n{new_email}"
            break


def remove_text_labels(doc: Document, labels_to_remove: Iterable[str]) -> None:
    """Remove phrases from the document wherever they appear."""

    for label in labels_to_remove:
        replace_text_in_paragraphs(doc.paragraphs, label, "")


def update_assistant_editors(doc: Document, remove_name: str) -> None:
    """Remove an entry from the assistant editors list."""

    for p in doc.paragraphs:
        if remove_name in p.text:
            elem = p._element
            parent = elem.getparent()
            parent.remove(elem)
            break


def insert_presidents_message(
    doc: Document, image_path: Path, message_text: str
) -> None:
    """Insert the president's message and optional image on page 3."""

    text = message_text if message_text else "<<Awaiting President's message>>"
    for i, p in enumerate(doc.paragraphs):
        if "President's Message" in p.text:
            if i + 1 < len(doc.paragraphs):
                target = doc.paragraphs[i + 1]
            else:
                target = doc.add_paragraph()
            target.text = text
            break


def extract_article_titles_from_toc(doc: Document) -> List[str]:
    """Return article titles listed under the ARTICLES section in the TOC."""
    toc_start = None
    paragraphs = doc.paragraphs
    # locate table of contents
    for i, p in enumerate(paragraphs):
        if "TABLE OF CONTENTS" in p.text.upper():
            toc_start = i
            break
    if toc_start is None:
        return []

    # find ARTICLES heading within TOC
    start = None
    for j in range(toc_start + 1, len(paragraphs)):
        text = paragraphs[j].text.strip()
        if text.upper().startswith("ARTICLES"):
            start = j + 1
            break
    if start is None:
        return []

    titles: List[str] = []
    import re

    for k in range(start, len(paragraphs)):
        line = paragraphs[k].text.strip()
        if not line:
            break
        if line.isupper():
            break
        match = re.match(r"(.+?)\.{2,}\d+$", line)
        if match:
            titles.append(match.group(1).strip())
        else:
            titles.append(line)

    return titles


def clear_articles(doc: Document):
    """Remove article sections based on TOC titles if available."""
    titles = extract_article_titles_from_toc(doc)

    def remove_range(start_p, end_p=None):
        body = doc.element.body
        end_el = end_p._element if end_p is not None else None
        started = False
        for el in list(body):
            if el is start_p._element:
                started = True
            if not started:
                continue
            if el is end_el:
                break
            tag = el.tag.rsplit("}", 1)[-1]
            if tag in ("p", "tbl"):
                body.remove(el)

    if titles:
        article_heading_idx = None

        # find start indices of each article title and detect the heading
        start_indices = []
        for title in titles:
            for i, p in enumerate(doc.paragraphs):
                if p.text.strip().upper() == title.upper():
                    start_indices.append(i)
                    if article_heading_idx is None and i > 0:
                        prev = doc.paragraphs[i - 1].text.strip().upper()
                        if prev == "ARTICLES":
                            article_heading_idx = i - 1
                    break
        start_indices.sort()
        if article_heading_idx is not None and (
            not start_indices or article_heading_idx < start_indices[0]
        ):
            start_indices.insert(0, article_heading_idx)

        start_paragraphs = [doc.paragraphs[i] for i in start_indices]
        for idx in reversed(range(len(start_paragraphs))):
            start_p = start_paragraphs[idx]
            end_p = (
                start_paragraphs[idx + 1] if idx + 1 < len(start_paragraphs) else None
            )
            remove_range(start_p, end_p)
        return

    # fallback to previous behaviour if we cannot parse TOC
    found = False
    start_idx = 0
    for i, p in enumerate(doc.paragraphs):
        if "ARTICLES" in p.text.upper():
            found = True
            start_idx = i
            break
    if found:
        remove_range(doc.paragraphs[start_idx])


def clear_articles_preserve_editorials(doc: Document) -> None:
    """Remove article content while keeping editorial sections."""

    headings = ["President's Message"]
    pages = map_pages_to_paragraphs(doc)
    para_index = {id(p): i for i, p in enumerate(doc.paragraphs)}

    last_editorial_page = 0
    for page_num, paragraphs in pages.items():
        for p in paragraphs:
            text = p.text.strip().lower()
            if text in [h.lower() for h in headings] or "editorial" in text:
                if page_num > last_editorial_page:
                    last_editorial_page = page_num

    if last_editorial_page == 0:
        clear_articles(doc)
        return

    article_start_idx = None
    for page in sorted(pages):
        if page >= last_editorial_page:
            for p in pages[page]:
                if p.text.strip().upper() == "ARTICLES":
                    article_start_idx = para_index.get(id(p))
                    break
            if article_start_idx is not None:
                break

    if article_start_idx is None:
        clear_articles(doc)
        return

    def remove_paragraph(paragraph):
        el = paragraph._element
        el.getparent().remove(el)

    for _ in range(len(doc.paragraphs) - article_start_idx):
        remove_paragraph(doc.paragraphs[article_start_idx])


def load_instructions(content_path: Path) -> dict:
    """Load instructions from ``instructions.json`` if present."""
    inst_file = content_path / "instructions.json"
    if inst_file.exists():
        try:
            with inst_file.open("r", encoding="utf-8") as f:
                data = json.load(f)
                if "format_front_and_footer" in data and not isinstance(
                    data["format_front_and_footer"], dict
                ):
                    data["format_front_and_footer"] = {}
                return data
        except Exception as e:
            print(f"Failed to read instructions: {e}")
    return {}


def delete_after_page(doc: Document, page_number: int) -> None:
    """Remove all paragraphs after ``page_number``.

    The function relies on :func:`map_pages_to_paragraphs` to determine the first
    paragraph of the next page and deletes everything that follows.
    """

    pages = map_pages_to_paragraphs(doc)

    # Find the first page that comes after ``page_number``
    next_page = None
    for num in sorted(pages):
        if num > page_number:
            next_page = num
            break

    if next_page is None:
        return

    paragraphs = pages.get(next_page)
    if not paragraphs:
        return

    first_el = paragraphs[0]._element
    idx = None
    for i, p in enumerate(doc.paragraphs):
        if p._element is first_el:
            idx = i
            break

    if idx is None:
        return

    while len(doc.paragraphs) > idx:
        el = doc.paragraphs[idx]._element
        el.getparent().remove(el)


def _find_last_editorial_page(doc: Document) -> Optional[int]:
    """Return the last page number containing an editorial heading."""

    pages = map_pages_to_paragraphs(doc)
    last = None
    for page_num, paragraphs in pages.items():
        for p in paragraphs:
            text = p.text.strip().lower()
            if "editorial" in text or "president's message" in text:
                if last is None or page_num > last:
                    last = page_num
    return last


def delete_after_editorial(doc: Document) -> None:
    """Remove all pages following the last editorial section."""

    last_page = _find_last_editorial_page(doc)
    if last_page is not None:
        delete_after_page(doc, last_page)


def _is_line_paragraph(paragraph) -> bool:
    """Return ``True`` if ``paragraph`` represents a horizontal line."""

    text = paragraph.text.strip()
    if text and len(text) >= 3 and all(ch in "-_—–" for ch in text):
        return True
    try:
        from docx.oxml.ns import qn
    except Exception:
        return False
    pPr = paragraph._p.pPr
    if pPr is not None:
        for b in pPr.findall(qn("w:pBdr")):
            bottom = b.find(qn("w:bottom"))
            if bottom is not None and bottom.get(qn("w:val")) != "nil":
                return True
    return False


def cleanup_black_lines(doc: Document) -> None:
    """Remove duplicate horizontal lines from each page."""

    pages = map_pages_to_paragraphs(doc)
    for paragraphs in pages.values():
        found = False
        for p in list(paragraphs):
            if _is_line_paragraph(p):
                if found:
                    el = p._element
                    el.getparent().remove(el)
                else:
                    found = True


def remove_pages_from(doc: Document, start_page: int) -> int:
    """Remove all pages beginning with ``start_page`` and return insertion index."""
    pages = map_pages_to_paragraphs(doc)
    paragraphs = pages.get(start_page)
    if not paragraphs:
        return len(doc.paragraphs)
    first_p = paragraphs[0]
    idx = doc.paragraphs.index(first_p)
    while len(doc.paragraphs) > idx:
        el = doc.paragraphs[idx]._element
        el.getparent().remove(el)
    return idx
  
def apply_basic_formatting(
    doc: Document, font_size: Optional[int], line_spacing: Optional[float]
) -> None:
    """Set font size and line spacing across all paragraphs."""
    for p in doc.paragraphs:
        if line_spacing is not None:
            p.paragraph_format.line_spacing = line_spacing
        for run in p.runs:
            if font_size is not None:
                run.font.size = Pt(font_size)


def append_article(doc: Document, article_doc: Document):
    for element in article_doc.element.body:
        doc.element.body.append(element)


def find_article_files(content_path: Path) -> List[Path]:
    """Return article files matching ``article*.docx`` case-insensitively."""
    pattern = "article"
    files = sorted(
        p
        for p in content_path.iterdir()
        if p.is_file()
        and p.name.lower().startswith(pattern)
        and p.suffix.lower() == ".docx"
    )
    if not files:
        logging.warning(
            "No article files found matching '%s*.docx' in %s", pattern, content_path
        )
    return files


def map_pages_to_paragraphs(doc: Document) -> Dict[int, List["Paragraph"]]:
    """Return a mapping of page numbers to paragraph objects.

    The detection relies on explicit ``w:br`` elements with ``w:type="page"``
    inserted by Word when a manual page break is present. If the document
    pagination depends solely on layout, the mapping may be inaccurate.
    """

    pages: Dict[int, List["Paragraph"]] = {1: []}
    current_page = 1
    for p in doc.paragraphs:
        pages.setdefault(current_page, []).append(p)
        if p._element.xpath('.//w:br[@w:type="page"]'):
            current_page += 1
            pages.setdefault(current_page, [])
    return pages


def autofit_first_table(doc: Document, page_num: int) -> None:
    """Autofit the first table on ``page_num`` if one exists."""

    pages = map_pages_to_paragraphs(doc)
    if page_num not in pages:
        return

    para_to_page = {p._element: n for n, ps in pages.items() for p in ps}

    try:
        from docx.table import Table
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:  # pragma: no cover - python-docx not installed
        Table = None  # type: ignore

    current_page = 1
    for el in doc.element.body.iterchildren():
        tag = el.tag.rsplit("}", 1)[-1]
        if tag == "p" and el in para_to_page:
            current_page = para_to_page[el]
            if el.xpath('.//w:br[@w:type="page"]'):
                current_page += 1
        elif tag == "tbl" and current_page == page_num and Table is not None:
            try:
                table = Table(el, doc)
                try:
                    table.autofit = True
                except Exception:
                    try:
                        tbl_pr = table._tbl.tblPr
                        if tbl_pr is None:
                            tbl_pr = OxmlElement("w:tblPr")
                            table._tbl.insert(0, tbl_pr)
                        layout = tbl_pr.find(qn("w:tblLayout"))
                        if layout is None:
                            layout = OxmlElement("w:tblLayout")
                            tbl_pr.append(layout)
                        layout.set(qn("w:type"), "autofit")
                        for col in table._tbl.findall(qn("w:gridCol")):
                            col.set(qn("w:w"), "0")
                    except Exception:
                        pass
            except Exception:
                pass
            break



def set_font_size(doc: Document, start_paragraph: int, size: int) -> None:
    """Apply ``size`` point font to paragraphs starting at ``start_paragraph``."""
    for p in doc.paragraphs[start_paragraph:]:
        for run in p.runs:
            run.font.size = Pt(size)


def set_line_spacing(doc: Document, start_paragraph: int, spacing: float) -> None:
    """Set line spacing for paragraphs starting at ``start_paragraph``."""
    for p in doc.paragraphs[start_paragraph:]:
        p.paragraph_format.line_spacing = spacing


def set_font_family(doc: Document, start_paragraph: int, font_name: str) -> None:
    """Set the font family for paragraphs starting at ``start_paragraph``."""
    for p in doc.paragraphs[start_paragraph:]:
        for run in p.runs:
            run.font.name = font_name


def format_front_and_footer(
    doc: Document,
    font_size: Optional[int] = None,
    line_spacing: Optional[float] = None,
) -> None:
    """Apply formatting to the front cover block and all footers."""

    # front cover paragraph usually contains "Volume" text
    for p in doc.paragraphs:
        if "Volume" in p.text:
            if line_spacing is not None:
                p.paragraph_format.line_spacing = line_spacing
            if font_size is not None:
                for run in p.runs:
                    run.font.size = Pt(font_size)
            break

    for section in doc.sections:
        for p in section.footer.paragraphs:
            if line_spacing is not None:
                p.paragraph_format.line_spacing = line_spacing
            if font_size is not None:
                for run in p.runs:
                    run.font.size = Pt(font_size)


def reuse_journal_page(doc: Document, source_doc: Document, page_number: int) -> None:
    """Copy the specified page from ``source_doc`` into ``doc``."""
    # Complex page-level manipulation is not implemented; placeholder only.
    pass


def reuse_info_for_authors(doc: Document, source_doc: Document) -> None:
    """Insert the 'Information for Authors' section from ``source_doc``."""
    pass


def reuse_membership_app(doc: Document, source_doc: Document) -> None:
    """Insert the membership application page from ``source_doc``."""
    pass


def add_editor_titles(doc: Document, page: int, editor_titles: List[str]) -> None:
    """Add editor titles under their images on the specified page."""
    pass


def remove_extra_spaces_in_author_line(
    doc: Document, page: int, article_index: int
) -> None:
    """Collapse multiple spaces in the given author line."""
    pass


def await_presidents_message_placeholder(doc: Document, page: int) -> None:
    """Insert a placeholder comment for the president's message."""
    pass


def update_author_line(
    doc: Document, page: int, old_name: str, new_name_with_creds: str
) -> None:
    """Replace the author line on the given page."""
    pass


def fix_orphaned_last_lines(doc: Document, page: int, column_index: int) -> None:
    """Ensure the last reference line wraps correctly."""
    pass


def convert_table_to_landscape(doc: Document, page: int, table_index: int) -> None:
    """Rotate a table on the specified page to landscape orientation."""
    pass


def move_paragraph_to_next_column(
    doc: Document, page: int, column_index: int, paragraph_text_match: str
) -> None:
    """Move a paragraph to the next column if it matches the text."""
    pass


def fix_apostrophe(
    doc: Document, page: int, search_word: str, correct_word: str
) -> None:
    """Replace a word with a corrected apostrophe."""
    replace_text_in_paragraphs(doc.paragraphs, search_word, correct_word)


def insert_line_space_before_subheading(
    doc: Document, page: int, subheading_list: Iterable[str]
) -> None:
    """Ensure a blank line before each subheading on the page."""
    pass


def move_section_to_next_column(
    doc: Document, page: int, column_index: int, section_heading: str
) -> None:
    """Move a section starting with heading to the next column."""
    pass


def insert_line_space_before_paragraph(
    doc: Document, page: int, preceding_text: str, new_paragraph_label: str
) -> None:
    """Insert a blank line before the paragraph that matches ``new_paragraph_label``."""
    pass


def indent_paragraph(
    doc: Document,
    page: int,
    column_index: int,
    paragraph_index: int,
    indent_width: float,
) -> None:
    """Apply left indent to a paragraph."""
    pass


def fix_orphaned_citation(doc: Document, page: int, orphan_text: str) -> None:
    """Merge an orphaned citation with the preceding line."""
    pass


def insert_line_space_before(doc: Document, page: int, subheading: str) -> None:
    """Ensure a blank line before the given subheading."""
    pass


def fix_separation_line_between_sections(
    doc: Document, page: int, line_shape_criteria
) -> None:
    """Ensure a solid line separates sections."""
    pass


def remove_unclear_text(doc: Document, page: int, search_pattern: str) -> None:
    """Remove or flag unclear text matching ``search_pattern``."""
    pass


def fix_page_numbering(doc: Document) -> None:
    """Renumber pages sequentially starting from the first section."""
    pass


def apply_hanging_indent_to_references(
    doc: Document,
    page_range_start: int,
    page_range_end: int,
    indent_width: float,
    line_spacing: float,
) -> None:
    """Apply hanging indent formatting to reference lists."""
    pass


def normalize_table_formatting(
    doc: Document,
    page: int,
    table_index: int,
    desired_font_size: int,
    desired_line_spacing: float,
) -> None:
    """Normalize fonts and spacing in tables."""
    pass


def detect_and_remove_extra_spaces(
    doc: Document, page_range: Iterable[int], pattern: str = "  "
) -> None:
    """Collapse multiple spaces across paragraphs in ``page_range``."""

    pages = map_pages_to_paragraphs(doc)
    for page_num in page_range:
        for p in pages.get(page_num, []):
            for run in p.runs:
                while pattern in run.text:
                    run.text = run.text.replace(pattern, " ")


def ensure_blank_line_before_headings(
    doc: Document, page_range: Iterable[int], heading_list: Iterable[str]
) -> None:
    """Ensure exactly one blank line precedes each heading."""
    pass


def split_long_paragraphs_across_columns(
    doc: Document, page: int, column_count: int
) -> None:
    """Attempt to split long paragraphs across columns."""
    pass


def update_table_of_contents(doc: Document) -> None:
    """Insert or refresh the table of contents."""
    pass


def apply_two_column_layout(doc: Document, start_page: int) -> None:
    """Set sections starting at ``start_page`` to use two text columns."""

    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return

    for idx, section in enumerate(doc.sections):
        if idx + 1 < start_page:
            continue

        # Prefer new API if available
        if hasattr(section, "text_columns"):
            try:
                section.text_columns.set_num(2)
                section.text_columns.spacing = Pt(36)  # ensure a small gap
            except Exception:
                pass
            continue

        sectPr = section._sectPr
        cols = sectPr.find(qn("w:cols"))
        if cols is None:
            cols = OxmlElement("w:cols")
            cols.set(qn("w:space"), "720")
            sectPr.append(cols)
        cols.set(qn("w:num"), "2")


def apply_page_borders(doc: Document, start_section: int, border_specs) -> None:
    """Apply borders to pages starting at ``start_section``.

    ``border_specs`` should be a mapping with keys for each border side
    (``"left"``, ``"right"``, ``"top"``, ``"bottom"``). Each side maps to a
    dictionary of border properties such as ``{"val": "single", "sz": 4}``.
    Only sides provided in ``border_specs`` are applied.
    """

    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        # If python-docx is unavailable we silently exit
        return

    valid_sides = {"left", "right", "top", "bottom"}

    for idx, section in enumerate(doc.sections):
        if idx < start_section:
            continue

        # Filter out unknown sides
        specs = {k: v for k, v in border_specs.items() if k in valid_sides}
        if not specs:
            continue

        # Use high level API when available
        if hasattr(section, "page_setup") and hasattr(
            section.page_setup, "left_border"
        ):
            ps = section.page_setup
            for side, spec in specs.items():
                try:
                    setattr(ps, f"{side}_border", spec)
                except Exception:
                    pass
            continue

        # Fallback to raw XML manipulation
        sectPr = section._sectPr
        for existing in sectPr.findall(qn("w:pgBorders")):
            sectPr.remove(existing)
        pgBorders = OxmlElement("w:pgBorders")
        pgBorders.set(qn("w:offsetFrom"), "text")

        for side, spec in specs.items():
            border = OxmlElement(f"w:{side}")
            for key, val in spec.items():
                border.set(qn(f"w:{key}"), str(val))
            pgBorders.append(border)
        sectPr.append(pgBorders)


def add_page_borders(doc: Document, start_section: int) -> None:
    """Add solid left and right borders for sections from ``start_section``."""

    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return

    for idx, section in enumerate(doc.sections):
        if idx < start_section:
            continue

        if hasattr(section, "page_setup") and hasattr(section.page_setup, "left_border"):  # type: ignore[attr-defined]
            try:
                ps = section.page_setup
                ps.left_border = ps.right_border = {
                    "val": "single",
                    "sz": 4,
                    "space": 0,
                    "color": "000000",
                }
                continue
            except Exception:
                pass

        sectPr = section._sectPr
        for existing in sectPr.findall(qn("w:pgBorders")):
            sectPr.remove(existing)
        pgBorders = OxmlElement("w:pgBorders")
        pgBorders.set(qn("w:offsetFrom"), "text")
        for side in ("left", "right"):
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")
            pgBorders.append(border)
        sectPr.append(pgBorders)


def apply_footer_layout(doc: Document, volume: str, issue: str, year: str) -> None:
    """Add standardized footers and leave the first page blank."""

    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return

    if not doc.sections:
        return

    first = doc.sections[0]
    first.different_first_page_header_footer = True

    def _clear_borders(table):
        tbl_pr = table._tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            table._tbl.insert(0, tbl_pr)
        borders = OxmlElement("w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "nil")
            borders.append(el)
        tbl_pr.append(borders)

    def _add_page_field(paragraph):
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        run._r.append(begin)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        run._r.append(instr)
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        run._r.append(separate)
        paragraph.add_run("1")
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        paragraph.add_run()._r.append(end)

    for idx, section in enumerate(doc.sections):
        if idx > 0:
            try:
                section.footer.is_linked_to_previous = False
            except Exception:
                pass
        footer = section.footer
        table = footer.add_table(rows=1, cols=3, width=section.page_width)
        _clear_borders(table)

        left_p = table.cell(0, 0).paragraphs[0]
        left_p.text = "The ABNFF Journal"
        left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in left_p.runs:
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0, 0, 0)

        center_cell = table.cell(0, 1)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "000000")
        center_cell._tc.get_or_add_tcPr().append(shd)
        center_p = center_cell.paragraphs[0]
        center_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_page_field(center_p)
        for r in center_p.runs:
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(255, 255, 255)

        right_p = table.cell(0, 2).paragraphs[0]
        right_p.text = f"Volume {volume} ({year}), Issue {issue}"
        right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for r in right_p.runs:
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0, 0, 0)


def validate_issue_number_and_volume(
    doc: Document, expected_volume: str, expected_issue: str, expected_year: str
) -> None:
    """Check volume/issue/year text appears once and matches expectations."""

    search = f"Volume {expected_volume}, Issue {expected_issue}"
    count_block = 0
    count_year = 0
    for p in doc.paragraphs:
        text = p.text
        if search in text:
            count_block += 1
        if expected_year in text:
            count_year += 1

    if count_block != 1 or count_year != 1:
        raise ValueError("Volume/issue/year text not found exactly once")


def save_pdf(doc_path: Path, pdf_path: Path) -> None:
    """Export ``doc_path`` to ``pdf_path`` using ``docx2pdf``.

    The underlying ``docx2pdf`` call relies on Microsoft Word. On some systems
    Word may report that the file is corrupted and abort the export. This error
    should not stop the rest of the update process, so we catch it and only
    log a warning instead of raising an exception.
    """

    try:
        from docx2pdf import convert

        convert(str(doc_path), str(pdf_path))
    except Exception as e:  # pragma: no cover - depends on Windows/Word
        err = str(e).lower()
        if "corrupted" in err:
            print(f"Warning: PDF export skipped—Word reported corruption: {e}")
        else:
            print(f"PDF export failed: {e}")


def update_journal(
    base_path: Path,
    content_path: Path,
    output_path: Path,
    volume: str,
    issue: str,
    month_year: str,
    cover_page_num: int = 1,
    start_page: Optional[int] = None,
    article_files: Optional[List[Path]] = None,
) -> None:
    """Run the update process and append ``article_files`` if provided.

    ``start_page`` specifies the page number where articles should be
    inserted. If ``None`` old articles are cleared automatically based on
    editorial headings.

    If ``article_files`` is ``None`` new articles are discovered using
    :func:`find_article_files` within ``content_path``.
    """
    doc = load_document(base_path)
    instructions = load_instructions(content_path)

    if "volume" in instructions:
        volume = str(instructions["volume"])
    if "issue" in instructions:
        issue = str(instructions["issue"])

    update_front_cover(doc, volume, issue, month_year, cover_page_num)
    apply_footer_layout(doc, volume, issue, month_year.split()[-1])
    update_business_information(
        doc,
        "2023",
        "Annual subscription rates are: institutions $550, individuals $220, and students $110",
    )
    header_text = f"Volume {volume}, Issue {issue}\n{month_year}"
    update_page2_header(doc, header_text, 2)
    pres_message_path = content_path / "president_message.txt"
    message_text = pres_message_path.read_text() if pres_message_path.exists() else ""
    insert_presidents_message(doc, content_path / "president.jpg", message_text)

    if start_page is not None:
        start_idx = remove_pages_from(doc, start_page)
        if start_idx == len(doc.paragraphs):
            clear_articles_preserve_editorials(doc)
            start_idx = len(doc.paragraphs)
    else:
        clear_articles_preserve_editorials(doc)
        start_idx = len(doc.paragraphs)
    files = (
        article_files if article_files is not None else find_article_files(content_path)
    )
    for article_file in files:
        article_doc = Document(article_file)
        append_article(doc, article_doc)
    if "font_size" in instructions:
        set_font_size(doc, start_idx, int(instructions["font_size"]))
    if "line_spacing" in instructions:
        set_line_spacing(doc, start_idx, float(instructions["line_spacing"]))
    if "font_family" in instructions:
        set_font_family(doc, start_idx, instructions["font_family"])
    if "delete_after_page" in instructions:
        try:
            delete_after_page(doc, int(instructions["delete_after_page"]))
        except Exception:
            pass
    if instructions.get("delete_after_editorial"):
        delete_after_editorial(doc)
    if instructions.get("cleanup_black_lines"):
        cleanup_black_lines(doc)
    if "autofit_table_on_page" in instructions:
        try:
            autofit_first_table(doc, int(instructions["autofit_table_on_page"]))
        except Exception:
            pass

    update_table_of_contents(doc)

    save_document(doc, output_path)
    pdf_path = output_path.with_suffix(".pdf")
    save_pdf(output_path, pdf_path)


def main_from_gui(
    base_doc: Path,
    content_folder: Path,
    output_doc: Path,
    volume: str,
    issue: str,
    month_year: str,
    cover_page_num: int = 1,
    start_page: Optional[int] = None,
    article_files: Optional[List[Path]] = None,
    font_size: Optional[int] = None,
    line_spacing: Optional[float] = None,
    font_family: Optional[str] = None,
) -> None:
    """Helper for GUI front-end."""
    inst_file = content_folder / "instructions.json"
    instructions = {}
    if inst_file.exists():
        try:
            with inst_file.open("r", encoding="utf-8") as f:
                instructions = json.load(f)
        except Exception:
            instructions = {}
    if font_size is not None:
        instructions["font_size"] = font_size
    if line_spacing is not None:
        instructions["line_spacing"] = line_spacing
    if font_family is not None:
        instructions["font_family"] = font_family
    if instructions:
        with inst_file.open("w", encoding="utf-8") as f:
            json.dump(instructions, f)

    update_journal(
        base_doc,
        content_folder,
        output_doc,
        volume,
        issue,
        month_year,
        cover_page_num,
        start_page,
        article_files,
    )


def main():
    parser = argparse.ArgumentParser(description="Update ABNFF Journal document")
    parser.add_argument("base_doc")
    parser.add_argument("content_folder")
    parser.add_argument("output_doc")
    parser.add_argument("--volume", required=True)
    parser.add_argument("--issue", required=True)
    parser.add_argument("--month-year", required=True, dest="month_year")
    parser.add_argument("--cover-page", type=int, default=1, dest="cover_page")
    parser.add_argument(
        "--start-page", type=int, default=None, dest="start_page",
        help="Page number where new articles begin"
    )
    args = parser.parse_args()

    base_path = Path(args.base_doc)
    content_path = Path(args.content_folder)
    output_path = (
        Path(args.output_doc)
        if args.output_doc
        else base_path.with_name(base_path.stem + "_updated.docx")
    )

    update_journal(
        base_path,
        content_path,
        output_path,
        args.volume,
        args.issue,
        args.month_year,
        args.cover_page,
        args.start_page,
        None,
    )


if __name__ == "__main__":
    main()
